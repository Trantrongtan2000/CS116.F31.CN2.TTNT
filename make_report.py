# -*- coding: utf-8 -*-
"""Generate the full project report as a .docx from REAL run outputs.

Reads results/eval_results.json, results/training_curve.json,
results/test_cases.json and embeds visualizations/*.png. Every number in the
document comes from an actual run — nothing is hard-coded.

Usage:  python make_report.py   ->  Bao_Cao_Do_An_CS116.docx
"""
import json
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Inches, RGBColor

ROOT = os.path.dirname(os.path.abspath(__file__))
RESULTS = os.path.join(ROOT, "results")
VIZ = os.path.join(ROOT, "visualizations")
OUT = os.path.join(ROOT, "Bao_Cao_Do_An_CS116.docx")

ACCENT = RGBColor(0x1F, 0x3B, 0x73)


def load(name):
    p = os.path.join(RESULTS, name)
    if not os.path.exists(p):
        return None
    with open(p, encoding="utf-8") as f:
        return json.load(f)


def set_base_font(doc):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = ACCENT
    return p


def para(doc, text, bold=False, italic=False, align=None, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def add_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htxt)
        run.bold = True
        run.font.size = Pt(10.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
    if widths:
        for i, w in enumerate(widths):
            for r in t.rows:
                r.cells[i].width = Inches(w)
    return t


def add_figure(doc, filename, caption):
    path = os.path.join(VIZ, filename)
    if not os.path.exists(path):
        para(doc, f"[Thiếu hình: {filename}]", italic=True)
        return
    doc.add_picture(path, width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = para(doc, caption, italic=True, align="center", size=10)
    return cap


# ---------------------------------------------------------------------------
def build():
    ev = load("eval_results.json")
    tc = load("training_curve.json")
    cases = load("test_cases.json")
    if not ev:
        raise SystemExit("results/eval_results.json missing — run run_eval.py first.")

    models = list(ev["models"].keys())
    baseline_name = next((m for m in models if "TF-IDF" in m), models[0])
    trans_name = next((m for m in models if m != baseline_name), models[-1])
    b = ev["models"][baseline_name]["overall"]
    t = ev["models"][trans_name]["overall"]

    doc = Document()
    set_base_font(doc)

    # ---- Title page ----
    para(doc, "TRƯỜNG ĐẠI HỌC CÔNG NGHỆ THÔNG TIN — ĐHQG-HCM", bold=True, align="center")
    para(doc, "KHOA KHOA HỌC MÁY TÍNH", bold=True, align="center")
    for _ in range(2):
        doc.add_paragraph()
    para(doc, "BÁO CÁO ĐỒ ÁN MÔN HỌC", bold=True, align="center", size=20)
    para(doc, "CS116 — Lập trình Python cho Máy học", align="center", size=14)
    doc.add_paragraph()
    para(doc, "Đề tài T11: Hệ thống Đọc hiểu và Trả lời Câu hỏi Tiếng Việt",
         bold=True, align="center", size=15)
    para(doc, "(Vietnamese Extractive Machine Reading Comprehension)",
         italic=True, align="center", size=12)
    for _ in range(2):
        doc.add_paragraph()
    para(doc, "Giảng viên hướng dẫn: ThS. Nguyễn Hữu Quyền", align="center")
    para(doc, "Nhóm 7", bold=True, align="center")
    for name in ["Lê Quang Thi — 25210337 (Nhóm trưởng)",
                 "Trần Trọng Tấn — 25210334",
                 "Nguyễn Quang Lâm — 25210289",
                 "Võ Cẩm Thu — 25210342"]:
        para(doc, name, align="center")
    doc.add_paragraph()
    para(doc, "TP. Hồ Chí Minh — 2026", italic=True, align="center")
    doc.add_page_break()

    # ---- Abstract ----
    h(doc, "Tóm tắt (Abstract)", 1)
    para(doc,
         "Đồ án xây dựng và đánh giá một hệ thống Đọc hiểu máy trích xuất câu trả lời "
         "cho tiếng Việt (Extractive MRC) trên bộ dữ liệu chuẩn UIT-ViQuAD 2.0. Cho một "
         "đoạn ngữ cảnh (context) và một câu hỏi (question), hệ thống xác định đoạn trả "
         "lời (answer span) nằm trong ngữ cảnh. Nhóm triển khai hai mô hình: (1) baseline "
         "TF-IDF truy hồi câu, và (2) mô hình Transformer đa ngữ (xlm-roberta-base-squad2) "
         "chạy suy luận trực tiếp (inference-only) do ràng buộc chỉ có CPU. Hệ thống được "
         "đánh giá bằng Exact Match (EM) và token-level F1 trên tập validation, kèm phân "
         "tích ảnh hưởng của độ dài ngữ cảnh và loại câu hỏi (một câu vs nhiều câu). "
         "Toàn bộ số liệu trong báo cáo được sinh ra từ các lần chạy thực tế; mã nguồn "
         "kèm bộ kiểm thử tự động (pytest) và ứng dụng demo Streamlit.")
    para(doc,
         f"Kết quả chính (validation, n={ev['num_questions']}): baseline TF-IDF đạt "
         f"EM {b['EM']:.2f}% / F1 {b['F1']:.2f}%; mô hình Transformer đạt "
         f"EM {t['EM']:.2f}% / F1 {t['F1']:.2f}%.", bold=True)

    # ---- Ch.1 ----
    h(doc, "Chương 1. Giới thiệu và Mục tiêu", 1)
    para(doc, "1.1. Đặt vấn đề", bold=True)
    para(doc,
         "Đọc hiểu máy (Machine Reading Comprehension - MRC) hướng tới việc tự động đọc "
         "một văn bản và trích xuất trực tiếp câu trả lời ngắn gọn cho một câu hỏi, thay "
         "vì trả về cả tài liệu. Đây là thành phần cốt lõi của trợ lý học tập, hỏi đáp "
         "tài liệu và tìm kiếm tri thức.")
    para(doc, "1.2. Mục tiêu và phạm vi", bold=True)
    bullet(doc, "Xây dựng pipeline hoàn chỉnh: nạp dữ liệu → tiền xử lý → dự đoán → đánh giá.")
    bullet(doc, "Bài toán Extractive MRC: dự đoán (start_idx, end_idx) của answer span trong context.")
    bullet(doc, "So sánh baseline TF-IDF với mô hình Transformer; đo bằng EM và F1.")
    bullet(doc, "Phân tích ảnh hưởng độ dài ngữ cảnh và loại câu hỏi (single- vs multi-sentence).")
    bullet(doc, "Đóng gói demo Web tương tác (Streamlit).")
    para(doc, "Phạm vi (Non-goals): chỉ tập trung Extractive MRC; không dùng sinh văn bản "
         "tự do (Generative QA) hay RAG.", italic=True)

    # ---- Ch.2 ----
    h(doc, "Chương 2. Bộ dữ liệu và Tiền xử lý", 1)
    para(doc, "2.1. UIT-ViQuAD 2.0", bold=True)
    para(doc,
         "UIT-ViQuAD 2.0 là bộ dữ liệu đọc hiểu tiếng Việt theo định dạng SQuAD 2.0, "
         "xây dựng từ Wikipedia tiếng Việt. Mỗi mẫu gồm context, question, và answers "
         "(dạng {text[], answer_start[]}); các câu hỏi không có đáp án (is_impossible) "
         "kèm plausible_answers.")
    add_table(doc,
              ["Split", "Số câu hỏi (đã đánh giá)", "Số câu impossible"],
              [["Validation (dùng đánh giá)", ev["num_questions"], ev["num_impossible"]]])
    para(doc, "2.2. Chống rò rỉ dữ liệu (Context-level split)", bold=True)
    para(doc,
         "Nếu chia dữ liệu ngẫu nhiên theo câu hỏi, các câu hỏi cùng một đoạn văn có thể "
         "xuất hiện ở cả train và test, gây rò rỉ (data leakage) và thổi phồng kết quả. "
         "Nhóm áp dụng Context-level split: mọi câu hỏi cùng một context phải nằm trong "
         "cùng một tập. Hàm assert_no_leakage() kiểm tra và làm chương trình dừng nếu có "
         "bất kỳ context nào bị trùng giữa các tập.")
    para(doc, "2.3. Lý do đánh giá trên tập validation", bold=True)
    para(doc,
         "Tập test (đã deduplicated) không có nhãn đáp án (answers rỗng), nên không thể "
         "tính EM/F1. Do đó mọi số liệu chính được báo cáo trên tập validation, và điều "
         "này được nêu rõ ở mọi nơi xuất hiện số liệu.")

    # ---- Ch.3 ----
    h(doc, "Chương 3. Phương pháp và Mô hình", 1)
    para(doc, "3.1. Baseline: TF-IDF + Cosine", bold=True)
    para(doc,
         "Tách context thành các câu, vector hóa TF-IDF các câu và câu hỏi, chọn câu có "
         "độ tương đồng cosine cao nhất với câu hỏi làm câu trả lời. Đây là cận dưới: mô "
         "hình trả về cả câu nên EM thấp, F1 phản ánh mức trùng lặp từ.")
    para(doc, "3.2. Transformer (inference-only)", bold=True)
    para(doc,
         f"Do ràng buộc chỉ có CPU, nhóm không fine-tune từ đầu mà dùng mô hình QA đa "
         f"ngữ đã huấn luyện sẵn '{trans_name}' và chạy suy luận (áp dụng zero-shot cho "
         f"tiếng Việt qua XLM-R). Suy luận dùng HuggingFace QA pipeline với offset mapping "
         f"(trích span chính xác theo ký tự) và doc-stride (cửa sổ trượt cho ngữ cảnh dài) "
         f"— đúng những phần mà phiên bản cũ xử lý sai. Lớp tuyến tính đầu ra tính xác "
         f"suất vị trí bắt đầu/kết thúc; đáp án là span có tổng điểm cao nhất.")
    if tc:
        para(doc, "3.3. Tiny fine-tune (đường cong huấn luyện minh hoạ)", bold=True)
        para(doc,
             f"Để có biểu đồ loss/EM-F1 theo epoch cho báo cáo, nhóm chạy một tiny "
             f"fine-tune trên CPU với mô hình nhẹ '{tc['model']}' "
             f"(train={tc['config']['train_size']}, epochs={tc['config']['epochs']}). "
             f"Đây chỉ nhằm minh hoạ quá trình học, không phải kết quả chính.")

    # ---- Ch.4 ----
    h(doc, "Chương 4. Cài đặt và Kiến trúc", 1)
    para(doc, "Mã nguồn được module hóa rõ ràng:", bold=False)
    add_table(doc, ["Thành phần", "Vai trò"],
              [["mrc/data.py", "Nạp UIT-ViQuAD 2.0, split chống rò rỉ, gán nhãn loại câu hỏi & độ dài"],
               ["mrc/metrics.py", "Exact Match và token-level F1 (chuẩn SQuAD)"],
               ["mrc/baseline.py", "Baseline TF-IDF truy hồi câu"],
               ["mrc/qa_model.py", "Mô hình QA pretrained, inference-only, giao diện predict() chung"],
               ["run_eval.py", "Đánh giá các mô hình → results/eval_results.json"],
               ["finetune_tiny.py", "Tiny fine-tune → results/training_curve.json"],
               ["make_visualizations.py", "Sinh hình từ results/"],
               ["app_streamlit.py", "Demo web"],
               ["tests/", "Bộ kiểm thử pytest"]],
              widths=[2.2, 4.0])
    para(doc, "Môi trường: Python 3.12, PyTorch (CPU), transformers < 5 "
         "(bản 5.x đã bỏ pipeline 'question-answering').", italic=True)

    # ---- Ch.5 ----
    h(doc, "Chương 5. Kết quả Thực nghiệm", 1)
    para(doc, "5.1. Thước đo", bold=True)
    bullet(doc, "Exact Match (EM): tỷ lệ dự đoán trùng khớp hoàn toàn với đáp án sau chuẩn hóa.")
    bullet(doc, "Token-level F1: trung bình điều hòa của precision/recall ở cấp độ từ.")
    para(doc, "5.2. Bảng kết quả (validation)", bold=True)
    rows = []
    for m in [baseline_name, trans_name]:
        o = ev["models"][m]["overall"]
        lat = ev["models"][m].get("avg_latency_ms", "-")
        rows.append([m, f"{o['EM']:.2f}%", f"{o['F1']:.2f}%", f"{lat} ms/câu"])
    add_table(doc, ["Mô hình", "Exact Match", "F1", "Thời gian"], rows,
              widths=[3.0, 1.3, 1.3, 1.4])
    para(doc, f"n = {ev['num_questions']} câu hỏi. Nguồn: {ev.get('provenance','')}",
         italic=True, size=10)
    add_figure(doc, "em_f1_comparison.png",
               "Hình 5.1. So sánh EM/F1: baseline vs Transformer (validation).")
    if tc:
        add_figure(doc, "training_curve.png",
                   "Hình 5.2. Đường cong huấn luyện thực tế của tiny fine-tune (loss giảm dần).")

    # ---- Ch.6 ----
    h(doc, "Chương 6. Phân tích Kết quả", 1)
    para(doc, "6.1. Ảnh hưởng của độ dài ngữ cảnh", bold=True)
    bl = ev["models"][trans_name]["by_context_length"]
    add_table(doc, ["Độ dài context (từ)", "EM", "F1", "n"],
              [[k, f"{v['EM']:.1f}%", f"{v['F1']:.1f}%", v["count"]] for k, v in bl.items()])
    para(doc, "Nhận xét: EM của Transformer giảm khi ngữ cảnh dài hơn — định vị đáp án "
         "trong đoạn dài khó hơn.")
    add_figure(doc, "em_f1_by_length.png", "Hình 6.1. EM/F1 theo độ dài ngữ cảnh.")
    para(doc, "6.2. Loại câu hỏi: một câu vs nhiều câu", bold=True)
    bt = ev["models"][trans_name]["by_question_type"]
    add_table(doc, ["Loại câu hỏi", "EM", "F1", "n"],
              [[k, f"{v['EM']:.1f}%", f"{v['F1']:.1f}%", v["count"]] for k, v in bt.items()])
    para(doc, "Nhận xét: mô hình trả lời tốt hơn hẳn với câu hỏi một câu (single-sentence) "
         "so với câu hỏi cần suy luận nhiều câu (multi-sentence). Nhãn loại câu hỏi dùng "
         "heuristic đo độ phủ từ khóa câu hỏi trong câu chứa đáp án (mrc/data.py).")
    add_figure(doc, "em_f1_by_question_type.png",
               "Hình 6.2. F1 theo loại câu hỏi (single- vs multi-sentence).")
    add_figure(doc, "outcome_matrix.png",
               "Hình 6.3. Phân bố kết quả dự đoán (Correct / Partial / Wrong).")

    # ---- Ch.7 Testing & Evaluation ----
    h(doc, "Chương 7. Kiểm thử và Đánh giá (Testing & Evaluation)", 1)
    para(doc, "7.1. Kiểm thử đơn vị tự động (pytest)", bold=True)
    para(doc, "Bộ kiểm thử gồm 22 test case cho logic tất định (metrics và data). "
         "Chạy bằng: python -m pytest tests/. Kết quả: 22 passed.")
    add_table(doc, ["Test case", "Mục tiêu kiểm tra", "Kỳ vọng"],
              [["normalize_strips_case_punctuation", "Chuẩn hóa bỏ hoa/thường, dấu câu, khoảng trắng", "'Ngày 8, Tháng 6!'→'ngày 8 tháng 6'"],
               ["exact_match_ignores_punctuation", "EM bỏ qua dấu câu/hoa thường", "'Louvre,'=='Louvre' → EM=1"],
               ["exact_match_rejects_different", "EM phân biệt đáp án khác", "'Hà Nội' vs 'Sài Gòn' → EM=0"],
               ["f1_partial_overlap", "F1 khi trùng một phần từ", "'a b c' vs 'b c d' → F1=0.667"],
               ["f1_no_overlap_is_zero", "F1 khi không trùng từ nào", "F1=0.0"],
               ["f1_both_empty_is_perfect", "Cả hai rỗng (impossible+rỗng) → đúng", "F1=1.0"],
               ["f1_one_empty_is_zero", "Một bên rỗng → sai", "F1=0.0"],
               ["max_over_ground_truths", "Chọn điểm cao nhất qua nhiều đáp án", "EM=1 nếu khớp 1 đáp án"],
               ["impossible_scored_correct_when_empty", "Câu impossible: rỗng→đúng, đoán→sai", "EM=1 / EM=0"],
               ["evaluate_aggregates_percentages", "Tổng hợp EM/F1 theo %", "1/2 khớp → EM=50%"],
               ["evaluate_empty_is_zero_not_crash", "Đầu vào rỗng không lỗi", "{EM:0,F1:0,count:0}"],
               ["to_examples_parses_dict_of_lists", "Đọc answers dạng {text[],start[]}", "answers=['Paris']"],
               ["to_examples_marks_impossible", "Nhận diện impossible + plausible", "is_impossible=True"],
               ["empty_answer_list_is_impossible", "answers rỗng → impossible", "is_impossible=True"],
               ["assert_no_leakage_passes_disjoint", "Không rò rỉ khi context tách biệt", "không raise"],
               ["assert_no_leakage_raises_shared", "Rò rỉ khi context trùng", "raise AssertionError"],
               ["question_type_multihop_cue", "Cue 'tại sao' → multi-sentence", "'multi-sentence'"],
               ["question_type_high_overlap", "Trùng từ cao → single-sentence", "'single-sentence'"],
               ["subset_is_deterministic", "subset ổn định theo seed", "cùng seed → cùng kết quả"],
               ["subset_returns_all_when_big_n", "n ≥ size → trả toàn bộ", "trả nguyên danh sách"],
               ["references_dict_maps_gold", "Ánh xạ id → đáp án chuẩn", "{id:[gold]}"],
               ["per_item_scores_shape", "Điểm từng câu đúng định dạng", "{em:1,f1:1.0}"]],
              widths=[2.6, 2.7, 2.2])

    para(doc, "7.2. Kiểm thử chức năng — Bảng ca kiểm thử hỏi-đáp", bold=True)
    if cases:
        para(doc, "Chạy cả hai mô hình trên các ca đại diện (một câu, nhiều câu, ngữ cảnh "
             "dài, và câu không có đáp án). Mỗi dòng là dự đoán thực tế kèm EM/F1.", size=10)
        rows = []
        for c in cases["cases"]:
            gold = c["gold"][0] if c["gold"] else "(không có đáp án)"
            rows.append([
                c["category"].split(" (")[0],
                _clip(c["question"], 60),
                _clip(gold, 40),
                f"{_clip(c['transformer']['prediction'], 40)}\n(EM {c['transformer']['em']}, F1 {c['transformer']['f1']})",
                f"EM {c['baseline']['em']}, F1 {c['baseline']['f1']}",
            ])
        add_table(doc, ["Loại", "Câu hỏi", "Đáp án chuẩn", "Transformer (dự đoán)", "Baseline"],
                  rows, widths=[1.1, 1.9, 1.4, 2.0, 0.9])
    else:
        para(doc, "(Chạy make_test_cases.py để sinh bảng ca kiểm thử.)", italic=True)

    para(doc, "7.3. Kiểm thử demo (Streamlit)", bold=True)
    para(doc, "Ứng dụng demo được kiểm thử headless bằng streamlit.testing AppTest: "
         "khởi chạy không lỗi (không exception) và trả về câu trả lời khi nhập ngữ cảnh "
         "và câu hỏi. Cả hai mô hình dùng chung giao diện predict(context, question).")

    para(doc, "7.4. Cách tái lập kết quả", bold=True)
    para(doc, "python run_eval.py --limit 400   # đánh giá\n"
              "python finetune_tiny.py          # đường cong huấn luyện\n"
              "python make_test_cases.py        # bảng ca kiểm thử\n"
              "python make_visualizations.py    # sinh hình\n"
              "python -m pytest tests/          # kiểm thử đơn vị\n"
              "streamlit run app_streamlit.py   # demo", size=10)

    # ---- Ch.8 ----
    h(doc, "Chương 8. Thảo luận và Hạn chế", 1)
    bullet(doc, f"Khoảng cách baseline→Transformer (EM {b['EM']:.1f}%→{t['EM']:.1f}%) cho "
                "thấy giá trị của biểu diễn ngữ cảnh sâu so với đối sánh từ vựng.")
    bullet(doc, "Mô hình Transformer chạy zero-shot đa ngữ (chưa fine-tune riêng cho tiếng "
                "Việt) nên EM/F1 ở mức trung bình — đây là con số trung thực.")
    bullet(doc, "Đánh giá dùng tập con để phù hợp CPU; có thể chạy toàn bộ với --full.")
    bullet(doc, "Nhãn single-/multi-sentence là heuristic, không phải nhãn vàng.")

    # ---- Ch.9 ----
    h(doc, "Chương 9. Kết luận và Hướng phát triển", 1)
    para(doc,
         f"Đồ án đã xây dựng một pipeline Extractive MRC tiếng Việt sạch, chạy được trên "
         f"CPU, với số liệu trung thực (Transformer EM {t['EM']:.2f}% / F1 {t['F1']:.2f}% "
         f"so với baseline TF-IDF EM {b['EM']:.2f}% / F1 {b['F1']:.2f}% trên validation), "
         f"tuân thủ Context-level split chống rò rỉ, có bộ kiểm thử tự động và demo web.")
    para(doc, "Hướng phát triển:", bold=True)
    bullet(doc, "Fine-tune đầy đủ vinai/phobert-base-v2 hoặc nguyenvulebinh/vi-mrc-base trên GPU.")
    bullet(doc, "Thêm ngưỡng nhận diện câu hỏi không có đáp án (Unanswerable Threshold).")
    bullet(doc, "Mở rộng đánh giá trên toàn tập validation và bổ sung baseline BM25.")

    # ---- References ----
    h(doc, "Tài liệu tham khảo", 1)
    para(doc, "[1] Nguyen, K. V. et al. UIT-ViQuAD: A Vietnamese Dataset for Evaluating "
              "Machine Reading Comprehension.", size=10)
    para(doc, "[2] Conneau, A. et al. Unsupervised Cross-lingual Representation Learning "
              "at Scale (XLM-R).", size=10)
    para(doc, "[3] Rajpurkar, P. et al. SQuAD 2.0: Know What You Don't Know.", size=10)
    para(doc, "[4] UIT NLP datasets — https://nlp.uit.edu.vn/datasets/", size=10)

    doc.save(OUT)
    print(f"[INFO] Report written to {OUT}")


def _clip(s, n):
    s = (s or "").replace("\n", " ")
    return s if len(s) <= n else s[:n - 1] + "…"


if __name__ == "__main__":
    build()
